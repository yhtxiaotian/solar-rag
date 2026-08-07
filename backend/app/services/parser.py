import re
from pathlib import Path

from app.core.config import settings
from app.services.chunker import ParsedBlock


class DocumentParseError(RuntimeError):
    pass


def _blocks_from_markdown(markdown: str) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    section: str | None = None
    buffer: list[str] = []

    def flush(kind: str = "paragraph") -> None:
        nonlocal buffer
        text = "\n".join(buffer).strip()
        if text:
            blocks.append(ParsedBlock(text, section=section, kind=kind))
        buffer = []

    for line in markdown.splitlines():
        if line.startswith("#"):
            flush()
            section = line.lstrip("#").strip()
            continue
        if line.startswith("|"):
            if buffer and not buffer[-1].startswith("|"):
                flush()
            buffer.append(line)
            continue
        if buffer and buffer[-1].startswith("|") and not line.startswith("|"):
            flush("table")
        if not line.strip():
            flush()
        else:
            buffer.append(line)
    flush("table" if buffer and buffer[0].startswith("|") else "paragraph")
    return blocks


def _parse_with_docling(path: Path) -> list[ParsedBlock]:
    from docling.document_converter import DocumentConverter

    result = DocumentConverter().convert(str(path))
    markdown = result.document.export_to_markdown()
    return _blocks_from_markdown(markdown)


def _ocr_page(page) -> str:  # type: ignore[no-untyped-def]
    if not settings.enable_ocr:
        return ""
    import numpy as np
    from rapidocr import RapidOCR

    pixmap = page.get_pixmap(dpi=200, alpha=False)
    image = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(pixmap.height, pixmap.width, pixmap.n)
    result = RapidOCR()(image)
    texts = getattr(result, "txts", None)
    if texts is None and isinstance(result, tuple) and result:
        lines = result[0] or []
        texts = [line[1] for line in lines if len(line) > 1]
    return "\n".join(texts or [])


def _parse_pdf(path: Path) -> list[ParsedBlock]:
    import fitz

    blocks: list[ParsedBlock] = []
    try:
        document = fitz.open(path)
        for page_number, page in enumerate(document, start=1):
            text = page.get_text("text", sort=True).strip()
            if len(re.sub(r"\s", "", text)) < 40:
                try:
                    text = _ocr_page(page).strip()
                except Exception as exc:  # OCR errors should identify the page.
                    raise DocumentParseError(f"第 {page_number} 页 OCR 失败：{exc}") from exc
            if text:
                blocks.append(ParsedBlock(text, page=page_number, section=f"第 {page_number} 页"))
        document.close()
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"PDF 解析失败：{exc}") from exc
    if not blocks:
        raise DocumentParseError("PDF 中没有可识别的正文")
    return blocks


def _parse_docx(path: Path) -> list[ParsedBlock]:
    from docx import Document as WordDocument

    doc = WordDocument(path)
    blocks: list[ParsedBlock] = []
    section: str | None = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            section = text
        else:
            blocks.append(ParsedBlock(text, section=section))
    for table in doc.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        widths = max(len(row) for row in rows)
        normalized = [row + [""] * (widths - len(row)) for row in rows]
        lines = ["| " + " | ".join(normalized[0]) + " |", "| " + " | ".join(["---"] * widths) + " |"]
        lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
        blocks.append(ParsedBlock("\n".join(lines), section=section, kind="table"))
    return blocks


def _parse_xlsx(path: Path) -> list[ParsedBlock]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ParsedBlock] = []
    for sheet in workbook.worksheets:
        rows = [["" if cell is None else str(cell).strip() for cell in row] for row in sheet.iter_rows(values_only=True)]
        rows = [row for row in rows if any(row)]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        rows = [row + [""] * (width - len(row)) for row in rows]
        header = rows[0]
        lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
        lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
        blocks.append(ParsedBlock("\n".join(lines), section=f"工作表：{sheet.title}", kind="table"))
    workbook.close()
    return blocks


def _parse_html(path: Path) -> list[ParsedBlock]:
    from bs4 import BeautifulSoup
    import trafilatura

    raw = path.read_text(encoding="utf-8", errors="replace")
    extracted = trafilatura.extract(raw, include_tables=True, output_format="markdown")
    if extracted:
        return _blocks_from_markdown(extracted)
    soup = BeautifulSoup(raw, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    return [ParsedBlock(soup.get_text("\n", strip=True))]


def parse_document(path: Path) -> list[ParsedBlock]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if settings.use_docling and suffix in {".docx", ".xlsx", ".html", ".htm"}:
        try:
            blocks = _parse_with_docling(path)
            if blocks:
                return blocks
        except Exception:
            pass
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".xlsx":
        return _parse_xlsx(path)
    if suffix in {".html", ".htm"}:
        return _parse_html(path)
    if suffix in {".txt", ".md"}:
        return _blocks_from_markdown(path.read_text(encoding="utf-8", errors="replace"))
    raise DocumentParseError(f"不支持的文件格式：{suffix}")
